import re
from docx import Document

def replace_text_in_doc(doc_path):
    doc = Document(doc_path)
    
    # We want to replace variations of the name
    replacements = {
        r'Loan Default Prediction & Bank Policy Underwriting': 'Loan Eligibility & Policy Underwriting Chatbot',
        r'Loan Default Prediction': 'Loan Eligibility',
        r'Default Prediction': 'Eligibility Assessment',
        r'default prediction': 'eligibility assessment',
        r'predict loan default probability': 'assess loan eligibility',
        r'reduced default rates': 'improved eligibility assessment accuracy'
    }

    # Function to replace text in a paragraph
    def replace_in_paragraph(p):
        for old, new in replacements.items():
            # simple run replacement can be tricky, so we join text, replace, and put in first run
            # clearing the rest if a match is found
            if re.search(old, p.text):
                full_text = p.text
                for key, val in replacements.items():
                    full_text = re.sub(key, val, full_text)
                
                # clear all runs
                for i in range(len(p.runs)):
                    p.runs[i].text = ""
                
                # put all text in the first run to preserve paragraph style
                if len(p.runs) > 0:
                    p.runs[0].text = full_text
                else:
                    p.add_run(full_text)

    # Search in paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p)
        
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)
                    
    doc.save(doc_path)
    print("Successfully updated the project name in the proposal document.")

if __name__ == "__main__":
    replace_text_in_doc(r'C:\Users\veera\OneDrive\Desktop\Project\Loan_AI_Chatbot_Proposal_Official.docx')
