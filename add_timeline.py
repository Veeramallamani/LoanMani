from docx import Document
from docx.shared import Pt, Inches

def add_timeline_to_proposal():
    doc_path = r'C:\Users\veera\OneDrive\Desktop\Project\Loan_AI_Chatbot_Proposal_Official.docx'
    
    try:
        doc = Document(doc_path)
        
        # Add a new heading for the timeline
        heading = doc.add_heading('Expected Project Timeline', level=1)
        
        # Add the timeline paragraph
        p = doc.add_paragraph(
            "The expected timeline for the completion and deployment of the 'Loan Default Prediction & Bank Policy Underwriting' "
            "web application is 45 days from the date of official acknowledgment and approval. This timeline encompasses all phases "
            "including backend model integration, frontend development, testing, and deployment."
        )
        
        # Add some bullet points for a breakdown if appropriate
        doc.add_paragraph("Phase 1 (Days 1-10): Requirements Gathering & UI/UX Design", style='List Bullet')
        doc.add_paragraph("Phase 2 (Days 11-25): Backend Development & ML Model Integration", style='List Bullet')
        doc.add_paragraph("Phase 3 (Days 26-35): Frontend Integration & Chatbot Development", style='List Bullet')
        doc.add_paragraph("Phase 4 (Days 36-45): Testing, Bug Fixing, & Final Deployment", style='List Bullet')
        
        doc.save(doc_path)
        print("Successfully added the expected timeline of 45 days to the document.")
        
    except Exception as e:
        print(f"Error modifying document: {e}")

if __name__ == "__main__":
    add_timeline_to_proposal()
